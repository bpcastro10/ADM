"""Lectura de CV desde archivos (PDF/DOCX/TXT)."""

from __future__ import annotations

from io import BytesIO
from typing import Tuple


def read_resume_bytes(filename: str, content: bytes) -> str:
    """
    Lee el contenido textual de un CV.
    Soporta: .pdf, .docx, .txt (UTF-8).
    """
    name = (filename or "").lower().strip()

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content))
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
            return "\n".join(t for t in texts if t).strip()
        except Exception as e:
            raise ValueError("No se pudo leer el PDF. Asegúrate de que no sea un PDF escaneado sin texto.") from e

    if name.endswith(".doc") and not name.endswith(".docx"):
        raise ValueError(
            "El formato .doc (Word antiguo) no está soportado. "
            "Convierta el archivo a .docx, .pdf o .txt."
        )

    if name.endswith(".docx"):
        try:
            import docx

            doc = docx.Document(BytesIO(content))
            parts = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            # tablas (a veces CVs usan tablas)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join((cell.text or "").strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts).strip()
        except Exception as e:
            raise ValueError("No se pudo leer el DOCX. Verifica que el archivo sea un .docx válido.") from e

    # default: txt
    try:
        return content.decode("utf-8").strip()
    except UnicodeDecodeError as e:
        raise ValueError("El archivo no es texto UTF-8 válido. Sube un .txt UTF-8, .pdf o .docx.") from e

